import docx
import pandas as pd
import os
location = "D:\\Shaheryar\\WORD"

finalData = pd.DataFrame()
files = os.listdir(location)
path = location+"\\"+files[0]
doc = docx.Document(path)
tb = doc.tables
rows = tb[0].rows
header = []
for i in range(1,39):
    if i==3:
        header.append("Longitude")
        continue
    if i==4:
        header.append("latitude")
        continue
        
    cols = rows[i].cells
    cell = cols [1]
    a= cell.text
    a.replace("\n","")
    header.append(a)
data = pd.DataFrame(columns = header)
for file in files:
    
    path = location+"\\"+file
    doc = docx.Document(path)
    tb = doc.tables
    rows = tb[0].rows
    Descrip = []

    for i in range(1,39):
      
        cols = rows[i].cells

        cell = cols [2]
        des= cell.text
        
        if i==3:
            Descrip.append(des[11:20])
            Descrip.append(des[30:])
            continue
        if i ==4:
            continue
        
        if i == 16:
            remCols = rows[i].cells
            remCell = cols[3]
            rem = remCell.text
            Descrip.append(des+" "+rem)
            continue
            
        Descrip.append(des)

    a_series = pd.Series(Descrip, index = data.columns)
    data = data.append(a_series, ignore_index=True)

data.to_excel(location+"\\"+"master_.xlsx")
